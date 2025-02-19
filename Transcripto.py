import sys
import os
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QVBoxLayout, QWidget, QPushButton,
    QLabel, QFileDialog, QLineEdit, QMessageBox
)
from PySide6.QtCore import Qt, QThread, Signal
from faster_whisper import WhisperModel
from moviepy import VideoFileClip
from docx import Document

def validate_and_convert_file(file_path):
    """
    Validate and convert the input file if needed (e.g., extract audio from video).
    """
    supported_audio_formats = ['.wav', '.mp3', '.m4a']
    supported_video_formats = ['.mp4', '.avi', '.mov', '.mkv']

    if not os.path.exists(file_path):
        raise FileNotFoundError("Le fichier n'existe pas.")

    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension in supported_audio_formats:
        return file_path
    elif file_extension in supported_video_formats:
        video = VideoFileClip(file_path)
        audio_path = os.path.splitext(file_path)[0] + '.mp3'
        video.audio.write_audiofile(audio_path)
        return audio_path
    else:
        raise ValueError("Format de fichier non pris en charge. Formats supportés : WAV, MP3, M4A, MP4, AVI, MOV, MKV.")

def transcribe_audio(file_path, output_path):
    """
    Transcribe audio using Faster-Whisper and save the result to a Word document.
    """
    model = WhisperModel("medium")
    segments, _ = model.transcribe(file_path)

    doc = Document()
    doc.add_heading("Transcription", level=1)
    
    for segment in segments:
        doc.add_paragraph(segment.text)
    
    doc.save(output_path)

class TranscriptionApp(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Transcription Audio/Video")
        self.setGeometry(300, 300, 500, 300)

        # Apply custom stylesheet
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f0f0f0;
            }

            QPushButton {
                background-color: #40a087;
                color: white;
                border: none;
                padding: 10px 20px;
                font-size: 14px;
                border-radius: 5px;
            }

            QPushButton:hover {
                background-color: #00805f;
            }

            QLabel {
                font-size: 14px;
                color: #333;
            }
            
            QMessageBox {
                background-color: #f0f0f0;
            }
        """)

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)

        self.layout = QVBoxLayout()
        self.central_widget.setLayout(self.layout)

        # Warning message at the top
        self.warning_label = QLabel(
            "Le temps de transcription peut être long.\n"
            "Veuillez désactiver la mise en veille automatique et éviter d'utiliser l'ordinateur."
        )
        self.warning_label.setStyleSheet("font-weight: bold; color: #d9534f;")
        self.warning_label.setAlignment(Qt.AlignCenter)
        self.layout.addWidget(self.warning_label)

        # File selection button
        self.file_label = QLabel("Aucun fichier sélectionné")
        self.file_label.setAlignment(Qt.AlignCenter)
        self.layout.addWidget(self.file_label)

        self.select_file_button = QPushButton("Sélectionner un fichier audio/vidéo")
        self.select_file_button.clicked.connect(self.select_file)
        self.layout.addWidget(self.select_file_button)

        # Output file selection button
        self.output_label = QLabel("Aucun emplacement de fichier de sortie sélectionné")
        self.output_label.setAlignment(Qt.AlignCenter)
        self.layout.addWidget(self.output_label)

        self.select_output_button = QPushButton("Sélectionner l'emplacement de sortie")
        self.select_output_button.clicked.connect(self.select_output_file)
        self.layout.addWidget(self.select_output_button)

        # Transcription button
        self.transcribe_button = QPushButton("Lancer la transcription")
        self.transcribe_button.clicked.connect(self.transcribe)
        self.layout.addWidget(self.transcribe_button)

        # Status label
        self.status_label = QLabel("")
        self.status_label.setAlignment(Qt.AlignCenter)
        self.layout.addWidget(self.status_label)

        self.selected_file = None
        self.output_file = None

    def select_file(self):
        """Open a file dialog to select an audio or video file."""
        file_dialog = QFileDialog(self)
        file_path, _ = file_dialog.getOpenFileName(
            self, "Sélectionner un fichier", "", "Audio/Video Files (*.wav *.mp3 *.m4a *.mp4 *.avi *.mov *.mkv)"
        )
        if file_path:
            self.selected_file = file_path
            self.file_label.setText(f"Fichier sélectionné : {os.path.basename(file_path)}")

    def select_output_file(self):
        """Open a file dialog to select the output Word file."""
        file_dialog = QFileDialog(self)
        file_path, _ = file_dialog.getSaveFileName(
            self, "Sélectionner l'emplacement de sortie", "", "Word Documents (*.docx)"
        )
        if file_path:
            if not file_path.endswith(".docx"):
                file_path += ".docx"
            self.output_file = file_path
            self.output_label.setText(f"Fichier de sortie : {os.path.basename(file_path)}")

    def transcribe(self):
        """Handle the transcription process."""
        if not self.selected_file:
            QMessageBox.warning(self, "Erreur", "Veuillez sélectionner un fichier avant de continuer.")
            return

        if not self.output_file:
            QMessageBox.warning(self, "Erreur", "Veuillez sélectionner un emplacement de fichier de sortie.")
            return

        try:
            # Validate and convert file if needed
            audio_path = validate_and_convert_file(self.selected_file)

            # Display temporary message
            self.status_label.setText("Transcription en cours... Veuillez patienter.")
            self.status_label.setStyleSheet("font-weight: bold;")
            QApplication.processEvents()

            # Transcribe and save
            transcribe_audio(audio_path, self.output_file)

            self.status_label.setText("Transcription terminée avec succès !")
            QMessageBox.information(self, "Succès", f"La transcription a été enregistrée dans {self.output_file}.")
        except Exception as e:
            self.status_label.setText("Une erreur s'est produite.")
            QMessageBox.critical(self, "Erreur", str(e))

if __name__ == "__main__":
    app = QApplication(sys.argv)

    window = TranscriptionApp()
    window.show()

    sys.exit(app.exec())